import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import json
import os
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
import undetected_chromedriver as uc
import time
import threading

class GPTTranslatorGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("GPT 文档翻译助手")
        self.root.geometry("800x700")

        # 根窗口自适应
        self.root.rowconfigure(0, weight=1)
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(1, weight=0)

        # 创建主框架
        self.main_frame = ttk.Frame(root, padding="10")
        self.main_frame.grid(row=0, column=0, sticky="nsew")
        self.main_frame.rowconfigure(0, weight=1)
        self.main_frame.columnconfigure(0, weight=1)

        # 创建界面元素
        self.create_widgets()

        # 状态栏
        self.status_var = tk.StringVar()
        self.status_var.set("就绪")
        self.status_bar = ttk.Label(root, textvariable=self.status_var, relief=tk.SUNKEN)
        self.status_bar.grid(row=1, column=0, sticky="ew")

        # 初始化变量
        self.driver = None
        self.paragraphs = []
        self.current_paragraph_index = 0
        self.translation_prompt = "请翻译成中文"
        self.paragraphs_per_batch = 1
        self.is_translating = False
        self.should_stop = False
        self.translation_thread = None

    def create_widgets(self):
        # 顶部控制区
        control_frame = ttk.LabelFrame(self.main_frame, text="控制面板", padding="5")
        control_frame.grid(row=0, column=0, sticky="ew", pady=5)
        control_frame.columnconfigure(1, weight=1)

        # 浏览器控制
        ttk.Button(control_frame, text="启动浏览器", command=self.toggle_browser).grid(row=0, column=0, padx=5, pady=5)

        # 文件选择
        ttk.Button(control_frame, text="选择Word文件", command=self.select_word_file).grid(row=0, column=1, padx=5, pady=5, sticky="w")

        # 段落数量显示
        self.paragraph_count_var = tk.StringVar(value="段落数量: 0")
        ttk.Label(control_frame, textvariable=self.paragraph_count_var).grid(row=0, column=2, padx=5, pady=5)

        # 段落选择区
        selection_frame = ttk.LabelFrame(self.main_frame, text="段落选择", padding="5")
        selection_frame.grid(row=1, column=0, sticky="ew", pady=5)
        selection_frame.columnconfigure(1, weight=1)

        # 选择模式
        self.selection_mode = tk.StringVar(value="all")
        ttk.Radiobutton(selection_frame, text="处理所有段落", variable=self.selection_mode, 
                       value="all", command=self.on_selection_mode_change).grid(row=0, column=0, padx=5, pady=5, sticky="w")
        ttk.Radiobutton(selection_frame, text="处理指定段落数", variable=self.selection_mode, 
                       value="custom", command=self.on_selection_mode_change).grid(row=0, column=1, padx=5, pady=5, sticky="w")
        ttk.Radiobutton(selection_frame, text="处理指定范围", variable=self.selection_mode, 
                       value="range", command=self.on_selection_mode_change).grid(row=0, column=2, padx=5, pady=5, sticky="w")

        # 指定段落数输入
        ttk.Label(selection_frame, text="处理段落数:").grid(row=1, column=0, padx=5, pady=5)
        self.custom_paragraph_count = tk.StringVar(value="10")
        self.custom_count_entry = ttk.Entry(selection_frame, textvariable=self.custom_paragraph_count, width=10)
        self.custom_count_entry.grid(row=1, column=1, padx=5, pady=5, sticky="w")

        # 段落范围输入
        range_frame = ttk.Frame(selection_frame)
        range_frame.grid(row=2, column=0, columnspan=3, sticky="ew", pady=5)
        ttk.Label(range_frame, text="起始段落:").grid(row=0, column=0, padx=5, pady=5)
        self.start_paragraph = tk.StringVar(value="1")
        self.start_entry = ttk.Entry(range_frame, textvariable=self.start_paragraph, width=8)
        self.start_entry.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(range_frame, text="结束段落:").grid(row=0, column=2, padx=5, pady=5)
        self.end_paragraph = tk.StringVar(value="10")
        self.end_entry = ttk.Entry(range_frame, textvariable=self.end_paragraph, width=8)
        self.end_entry.grid(row=0, column=3, padx=5, pady=5)

        # 参数设置区
        settings_frame = ttk.LabelFrame(self.main_frame, text="参数设置", padding="5")
        settings_frame.grid(row=2, column=0, sticky="ew", pady=5)
        settings_frame.columnconfigure(1, weight=1)

        # 每批段落数
        ttk.Label(settings_frame, text="每批段落数:").grid(row=0, column=0, padx=5, pady=5)
        self.batch_size_var = tk.StringVar(value="1")
        ttk.Entry(settings_frame, textvariable=self.batch_size_var, width=10).grid(row=0, column=1, padx=5, pady=5, sticky="w")

        # 翻译提示
        ttk.Label(settings_frame, text="翻译提示:").grid(row=1, column=0, padx=5, pady=5)
        self.prompt_var = tk.StringVar(value="请翻译成中文")
        ttk.Entry(settings_frame, textvariable=self.prompt_var, width=50).grid(row=1, column=1, padx=5, pady=5, sticky="ew")

        # 翻译进度区
        progress_frame = ttk.LabelFrame(self.main_frame, text="翻译进度", padding="5")
        progress_frame.grid(row=3, column=0, sticky="nsew", pady=5)
        progress_frame.rowconfigure(0, weight=1)
        progress_frame.columnconfigure(0, weight=1)

        # 进度显示
        self.progress_text = tk.Text(progress_frame, height=10, wrap="word")
        self.progress_text.grid(row=0, column=0, sticky="nsew")
        progress_scroll = ttk.Scrollbar(progress_frame, orient="vertical", command=self.progress_text.yview)
        progress_scroll.grid(row=0, column=1, sticky="ns")
        self.progress_text['yscrollcommand'] = progress_scroll.set

        # 底部按钮区
        button_frame = ttk.Frame(self.main_frame)
        button_frame.grid(row=4, column=0, pady=5)
        self.start_button = ttk.Button(button_frame, text="开始翻译", command=self.start_translation)
        self.start_button.grid(row=0, column=0, padx=5)
        self.stop_button = ttk.Button(button_frame, text="停止翻译", command=self.stop_translation, state="disabled")
        self.stop_button.grid(row=0, column=1, padx=5)
        ttk.Button(button_frame, text="导出结果", command=self.show_export_menu).grid(row=0, column=2, padx=5)

    def on_selection_mode_change(self):
        """当选择模式改变时更新界面"""
        mode = self.selection_mode.get()
        if mode == "all":
            self.custom_count_entry.config(state="disabled")
            self.start_entry.config(state="disabled")
            self.end_entry.config(state="disabled")
        elif mode == "custom":
            self.custom_count_entry.config(state="normal")
            self.start_entry.config(state="disabled")
            self.end_entry.config(state="disabled")
        elif mode == "range":
            self.custom_count_entry.config(state="disabled")
            self.start_entry.config(state="normal")
            self.end_entry.config(state="normal")

    def toggle_browser(self):
        if self.driver is None:
            self.start_browser()
        else:
            self.stop_browser()

    def start_browser(self):
        try:
            options = uc.ChromeOptions()
            driver = uc.Chrome(version_main=136, options=options)
            self.driver = driver
            self.driver.get("https://chat.openai.com/")
            self.driver.maximize_window()
            self.status_var.set("浏览器已启动，请登录")

            # 登录检测
            try:
                WebDriverWait(self.driver, 20).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, "button.composer-btn"))
                )
                self.status_var.set("已经登录成功")
            except TimeoutException:
                pass

        except Exception as e:
            messagebox.showerror("错误", f"启动浏览器失败: {e}")

    def stop_browser(self):
        if self.driver:
            self.driver.quit()
            self.driver = None
            self.status_var.set("浏览器已关闭")

    def select_word_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        if file_path:
            try:
                doc = Document(file_path)
                self.paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                self.paragraph_count_var.set(f"段落数量: {len(self.paragraphs)}")
                self.current_paragraph_index = 0
                self.status_var.set(f"已加载文件: {os.path.basename(file_path)}")
                
                # 更新自定义段落数的最大值
                if self.paragraphs:
                    self.custom_paragraph_count.set(str(min(10, len(self.paragraphs))))
                    # 更新范围输入框的默认值
                    self.start_paragraph.set("1")
                    self.end_paragraph.set(str(min(10, len(self.paragraphs))))
            except Exception as e:
                messagebox.showerror("错误", f"读取文件失败: {e}")

    def start_translation(self):
        if not self.driver:
            messagebox.showwarning("警告", "请先启动浏览器！")
            return
        if not self.paragraphs:
            messagebox.showwarning("警告", "请先选择Word文件！")
            return

        try:
            self.paragraphs_per_batch = int(self.batch_size_var.get())
            self.translation_prompt = self.prompt_var.get()
            
            # 确定要处理的段落范围
            if self.selection_mode.get() == "all":
                self.target_paragraph_count = len(self.paragraphs)
                self.current_paragraph_index = 0
            elif self.selection_mode.get() == "custom":
                custom_count = int(self.custom_paragraph_count.get())
                if custom_count <= 0 or custom_count > len(self.paragraphs):
                    messagebox.showerror("错误", f"段落数必须在1到{len(self.paragraphs)}之间！")
                    return
                self.target_paragraph_count = custom_count
                self.current_paragraph_index = 0
            elif self.selection_mode.get() == "range":
                start_idx = int(self.start_paragraph.get()) - 1  # 转换为0基索引
                end_idx = int(self.end_paragraph.get())
                if start_idx < 0 or end_idx > len(self.paragraphs) or start_idx >= end_idx:
                    messagebox.showerror("错误", f"段落范围必须在1到{len(self.paragraphs)}之间，且起始段落必须小于结束段落！")
                    return
                self.target_paragraph_count = end_idx - start_idx
                self.current_paragraph_index = start_idx
                self.range_start_index = start_idx  # 保存范围起始索引
            
            # 重置状态
            self.is_translating = True
            self.should_stop = False
            
            # 更新按钮状态
            self.start_button.config(state="disabled")
            self.stop_button.config(state="normal")
            
            # 清空进度显示
            self.progress_text.delete("1.0", tk.END)
            
            # 在新线程中开始翻译
            self.translation_thread = threading.Thread(target=self.translation_worker)
            self.translation_thread.daemon = True
            self.translation_thread.start()
            
        except ValueError:
            messagebox.showerror("错误", "每批段落数必须是数字！")

    def stop_translation(self):
        """停止翻译过程"""
        self.should_stop = True
        self.is_translating = False
        self.status_var.set("正在停止翻译...")
        
        # 更新按钮状态
        self.start_button.config(state="normal")
        self.stop_button.config(state="disabled")

    def translation_worker(self):
        """翻译工作线程"""
        try:
            # 计算结束索引
            if self.selection_mode.get() == "range":
                end_index = self.range_start_index + self.target_paragraph_count
            else:
                end_index = self.target_paragraph_count
            
            while self.current_paragraph_index < end_index and not self.should_stop:
                # 获取下一批段落
                batch_end_index = min(self.current_paragraph_index + self.paragraphs_per_batch, end_index)
                batch = self.paragraphs[self.current_paragraph_index:batch_end_index]
                
                # 构建消息
                message = f"{self.translation_prompt}\n\n" + "\n\n".join(batch)
                
                # 发送消息
                if self.send_message_to_chatgpt(message):
                    # 计算显示的段落号（从1开始）
                    display_start = self.current_paragraph_index + 1
                    display_end = batch_end_index
                    
                    # 在主线程中更新状态
                    self.root.after(0, lambda: self.status_var.set(f"正在翻译第 {display_start} 到 {display_end} 段..."))
                    
                    response = self.get_last_response()
                    if response and not self.should_stop:
                        # 在主线程中更新进度显示
                        self.root.after(0, lambda: self.update_progress_display(batch, response, display_start, display_end))
                        
                        # 更新索引
                        self.current_paragraph_index = batch_end_index
                        
                        # 短暂延迟
                        time.sleep(1)
                    else:
                        if self.should_stop:
                            break
                        else:
                            self.root.after(0, lambda: self.status_var.set("获取翻译结果失败"))
                            break
                else:
                    if self.should_stop:
                        break
                    else:
                        self.root.after(0, lambda: self.status_var.set("发送消息失败"))
                        break
            
            # 翻译完成或停止
            if self.should_stop:
                self.root.after(0, lambda: self.status_var.set("翻译已停止"))
            else:
                self.root.after(0, lambda: self.status_var.set("翻译完成"))
                
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("错误", f"翻译过程中出错: {e}"))
        finally:
            # 恢复按钮状态
            self.root.after(0, lambda: self.start_button.config(state="normal"))
            self.root.after(0, lambda: self.stop_button.config(state="disabled"))
            self.is_translating = False

    def update_progress_display(self, batch, response, start_idx, end_idx):
        """更新进度显示"""
        self.progress_text.insert(tk.END, f"原文 ({start_idx}-{end_idx}):\n")
        for p in batch:
            self.progress_text.insert(tk.END, f"{p}\n")
        self.progress_text.insert(tk.END, f"\n翻译结果:\n{response}\n\n")
        self.progress_text.see(tk.END)

    def send_message_to_chatgpt(self, message):
        try:
            selectors = [
                "p[data-placeholder='询问任何问题']",
                "p[data-placeholder='Ask anything']",
                "div[contenteditable='true']",
                "div[role='textbox']"
            ]
            
            textarea = None
            for selector in selectors:
                textarea = self.wait_for_element(By.CSS_SELECTOR, selector)
                if textarea:
                    break
            
            if textarea:
                self.driver.execute_script("arguments[0].innerHTML = arguments[1]", textarea, message)
                time.sleep(1)
                
                send_button = self.wait_for_element(By.CSS_SELECTOR, "button[data-testid='send-button'], button[aria-label='发送消息']")
                if send_button:
                    send_button.click()
                    return True
            return False
        except Exception as e:
            print(f"发送消息时出错: {str(e)}")
            return False

    def wait_for_element(self, by, value, timeout=10):
        try:
            element = WebDriverWait(self.driver, timeout).until(
                EC.presence_of_element_located((by, value))
            )
            return element
        except TimeoutException:
            return None

    def get_last_response(self):
        try:
            max_wait_time = 60
            start_time = time.time()
            last_text = ""
            while time.time() - start_time < max_wait_time and not self.should_stop:
                try:
                    articles = self.driver.find_elements(By.CSS_SELECTOR, "article[data-testid^='conversation-turn-']")
                    if not articles:
                        time.sleep(1)
                        continue
                    last_article = articles[-1]
                    copy_button = None
                    try:
                        copy_button = last_article.find_element(By.CSS_SELECTOR, "button[data-testid='copy-turn-action-button']")
                    except Exception:
                        copy_button = None
                    markdown_div = None
                    try:
                        markdown_div = last_article.find_element(By.CSS_SELECTOR, "div.markdown.prose")
                    except Exception:
                        markdown_div = None
                    if copy_button and copy_button.is_displayed() and markdown_div:
                        current_text = markdown_div.text
                        if current_text and current_text != last_text:
                            return current_text
                        last_text = current_text
                    time.sleep(1)
                except Exception as e:
                    print(f"查找回复时出错: {str(e)}")
                    time.sleep(1)
            return None
        except Exception as e:
            print(f"获取回复时出错: {str(e)}")
            return None

    def show_export_menu(self):
        """显示导出选项菜单"""
        if not self.progress_text.get("1.0", tk.END).strip():
            messagebox.showwarning("警告", "没有可导出的翻译结果！")
            return

        # 创建导出选项窗口
        export_window = tk.Toplevel(self.root)
        export_window.title("导出选项")
        export_window.geometry("300x150")
        export_window.transient(self.root)
        export_window.grab_set()

        # 内容选择
        content_var = tk.StringVar(value="all")
        ttk.Radiobutton(export_window, text="导出全部内容（原文+翻译）", variable=content_var, value="all").pack(pady=5, padx=10, anchor="w")
        ttk.Radiobutton(export_window, text="仅导出翻译结果", variable=content_var, value="translation").pack(pady=5, padx=10, anchor="w")

        # 格式选择
        format_var = tk.StringVar(value="txt")
        ttk.Radiobutton(export_window, text="导出为TXT文件", variable=format_var, value="txt").pack(pady=5, padx=10, anchor="w")
        ttk.Radiobutton(export_window, text="导出为Word文件", variable=format_var, value="word").pack(pady=5, padx=10, anchor="w")

        # 确认按钮
        ttk.Button(export_window, text="确认导出", 
                  command=lambda: self.export_results(content_var.get(), format_var.get(), export_window)).pack(pady=10)

    def export_results(self, content_type="all", format_type="txt", export_window=None):
        """导出结果到文件
        content_type: 'all' 或 'translation'
        format_type: 'txt' 或 'word'
        """
        if not self.progress_text.get("1.0", tk.END).strip():
            messagebox.showwarning("警告", "没有可导出的翻译结果！")
            return

        # 获取文件扩展名
        ext = ".docx" if format_type == "word" else ".txt"
        file_types = [("Word files", "*.docx")] if format_type == "word" else [("Text files", "*.txt"), ("All files", "*.*")]

        # 获取保存路径
        file_path = filedialog.asksaveasfilename(
            defaultextension=ext,
            filetypes=file_types
        )
        
        if not file_path:
            return

        try:
            # 获取文本内容
            full_text = self.progress_text.get("1.0", tk.END)
            
            if content_type == "translation":
                # 提取翻译结果
                translation_lines = []
                lines = full_text.split('\n')
                is_translation = False
                for line in lines:
                    if line.startswith("翻译结果:"):
                        is_translation = True
                        continue
                    elif line.startswith("原文 ("):
                        is_translation = False
                        continue
                    if is_translation and line.strip():
                        translation_lines.append(line)
                content = '\n'.join(translation_lines)
            else:
                content = full_text

            # 根据格式保存
            if format_type == "word":
                doc = Document()
                # 按段落分割并添加到Word文档
                paragraphs = content.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
                doc.save(file_path)
            else:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(content)

            messagebox.showinfo("成功", f"已导出到: {file_path}")
            
            # 关闭导出选项窗口
            if export_window:
                export_window.destroy()

        except Exception as e:
            messagebox.showerror("错误", f"导出失败: {e}")

def main():
    root = tk.Tk()
    app = GPTTranslatorGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main() 