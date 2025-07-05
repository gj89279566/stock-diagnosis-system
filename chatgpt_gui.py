import tkinter as tk
from tkinter import ttk, messagebox
import json
import os
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from webdriver_manager.chrome import ChromeDriverManager
import undetected_chromedriver as uc
import time
# 新增导出相关库
import html2text
import pdfkit
# 新增HTML解析和代码高亮相关库
from bs4 import BeautifulSoup
import re
from pygments import highlight
from pygments.lexers import get_lexer_by_name
from pygments.formatters import get_formatter_by_name
import tkinter.font as tkfont

class ChatGPTGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("ChatGPT 自动化助手")
        self.root.geometry("800x600")

        # 根窗口自适应
        self.root.rowconfigure(0, weight=1)
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(1, weight=0)

        # 创建主框架
        self.main_frame = ttk.Frame(root, padding="10")
        self.main_frame.grid(row=0, column=0, sticky="nsew")
        self.main_frame.rowconfigure(0, weight=1)
        self.main_frame.columnconfigure(0, weight=1)
        # 聊天区域
        self.create_chat_section()

        # 聊天轮次数据
        self.turn_data = []  # 每一轮为 {"question": ..., "answer": ...}

        # 状态栏
        self.status_var = tk.StringVar()
        self.status_var.set("就绪")
        self.status_bar = ttk.Label(root, textvariable=self.status_var, relief=tk.SUNKEN)
        self.status_bar.grid(row=1, column=0, sticky="ew")

        # 初始化浏览器
        self.driver = None
        self.is_first_message = True

        # 初始化代码高亮样式
        self.init_code_highlighting()

    def init_code_highlighting(self):
        """初始化代码高亮样式"""
        # 创建代码高亮样式
        self.code_style = {
            'python': {'background': '#f8f8f8', 'foreground': '#000000'},
            'javascript': {'background': '#f8f8f8', 'foreground': '#000000'},
            'html': {'background': '#f8f8f8', 'foreground': '#000000'},
            'css': {'background': '#f8f8f8', 'foreground': '#000000'},
            'default': {'background': '#f8f8f8', 'foreground': '#000000'}
        }
        
        # 为每种语言创建标签
        for lang, style in self.code_style.items():
            self.gpt_history.tag_configure(f"code_{lang}", 
                                        font=("Courier New", 12),
                                        background=style['background'],
                                        foreground=style['foreground'],
                                        spacing1=5,  # 代码块上方间距
                                        spacing3=5)  # 代码块下方间距

    def process_html_content(self, html_content):
        """处理HTML内容，转换为富文本格式"""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 处理代码块
        for code_block in soup.find_all('pre'):
            code = code_block.get_text()
            # 尝试检测语言
            lang = 'default'
            if code_block.get('class'):
                for cls in code_block['class']:
                    if cls.startswith('language-'):
                        lang = cls.replace('language-', '')
                        break
            
            # 应用代码高亮
            try:
                lexer = get_lexer_by_name(lang)
                formatter = get_formatter_by_name('html')
                highlighted_code = highlight(code, lexer, formatter)
                code_block.replace_with(BeautifulSoup(highlighted_code, 'html.parser'))
            except:
                # 如果高亮失败，使用普通代码格式
                code_block.replace_with(BeautifulSoup(f'<code class="code_{lang}">{code}</code>', 'html.parser'))
        
        # 处理其他HTML元素
        for tag in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li', 'table', 'tr', 'td', 'th']):
            if tag.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag.name[1])
                tag['style'] = f'font-size: {20 - level * 2}pt; font-weight: bold;'
            elif tag.name == 'ul':
                tag['style'] = 'list-style-type: disc; margin-left: 20px;'
            elif tag.name == 'ol':
                tag['style'] = 'list-style-type: decimal; margin-left: 20px;'
            elif tag.name == 'table':
                tag['style'] = 'border-collapse: collapse; margin: 10px 0;'
            elif tag.name in ['td', 'th']:
                tag['style'] = 'border: 1px solid #ccc; padding: 5px;'
        
        return str(soup)

    def create_chat_section(self):
        # 聊天框架
        chat_frame = ttk.LabelFrame(self.main_frame, text="聊天", padding="5")
        chat_frame.grid(row=0, column=0, sticky="nsew", pady=5)
        chat_frame.rowconfigure(2, weight=1)
        chat_frame.columnconfigure(0, weight=1)
        chat_frame.columnconfigure(1, weight=1)
        # 顶部输入区
        input_frame = ttk.Frame(chat_frame)
        input_frame.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(0, 8))
        input_frame.columnconfigure(0, weight=1)
        input_frame.columnconfigure(1, weight=0)
        input_frame.columnconfigure(2, weight=1)
        self.message_text = tk.Text(input_frame, height=3, font=("Arial", 14))
        self.message_text.grid(row=0, column=0, columnspan=3, sticky="ew")
        # 第二行：按钮横排
        send_btn = ttk.Button(input_frame, text="发送", command=self.send_message)
        send_btn.grid(row=1, column=0, padx=(0, 0), sticky="w")
        self.start_button = ttk.Button(input_frame, text="启动浏览器", command=self.toggle_browser)
        self.start_button.grid(row=1, column=1, padx=(8, 0), sticky="w")
        # 预留第三列用于拉伸
        # input_frame.columnconfigure(2, weight=1) # 已有

        # 历史区采用 PanedWindow 横向分栏
        history_paned = tk.PanedWindow(chat_frame, orient=tk.HORIZONTAL, sashrelief=tk.RAISED, showhandle=True)
        history_paned.grid(row=2, column=0, columnspan=2, sticky="nsew", pady=(0, 8))
        # PanedWindow自适应
        chat_frame.rowconfigure(2, weight=1)
        chat_frame.columnconfigure(0, weight=1)
        chat_frame.columnconfigure(1, weight=1)

        # 左侧用户历史 Frame，设置初始宽度
        user_hist_outer = ttk.Frame(history_paned, width=350)
        user_hist_outer.rowconfigure(1, weight=1)
        user_hist_outer.columnconfigure(0, weight=1)
        user_hist_label = ttk.Label(user_hist_outer, text="用户历史", font=("Arial", 12, "bold"))
        user_hist_label.grid(row=0, column=0, sticky="w", padx=(0, 0), pady=(0, 2))
        user_hist_content_frame = ttk.Frame(user_hist_outer)
        user_hist_content_frame.grid(row=1, column=0, sticky="nsew")
        user_hist_content_frame.rowconfigure(0, weight=1)
        user_hist_content_frame.columnconfigure(0, weight=1)
        self.user_history = tk.Text(user_hist_content_frame, height=14, font=("Arial", 14), wrap="word")
        self.user_history.grid(row=0, column=0, sticky="nsew")
        self.user_history.config(state=tk.DISABLED)
        user_scroll = ttk.Scrollbar(user_hist_content_frame, orient="vertical", command=self.user_history.yview)
        user_scroll.grid(row=0, column=1, sticky="ns")
        self.user_history['yscrollcommand'] = user_scroll.set

        # 右侧GPT历史 Frame，设置初始宽度
        gpt_hist_outer = ttk.Frame(history_paned, width=350)
        gpt_hist_outer.rowconfigure(1, weight=1)
        gpt_hist_outer.columnconfigure(0, weight=1)
        gpt_hist_label = ttk.Label(gpt_hist_outer, text="GPT历史", font=("Arial", 12, "bold"))
        gpt_hist_label.grid(row=0, column=0, sticky="w", padx=(0, 0), pady=(0, 2))
        gpt_hist_content_frame = ttk.Frame(gpt_hist_outer)
        gpt_hist_content_frame.grid(row=1, column=0, sticky="nsew")
        gpt_hist_content_frame.rowconfigure(0, weight=1)
        gpt_hist_content_frame.columnconfigure(0, weight=1)
        
        # 使用Text控件并配置富文本支持
        self.gpt_history = tk.Text(gpt_hist_content_frame, height=14, font=("Arial", 14), wrap="word")
        self.gpt_history.grid(row=0, column=0, sticky="nsew")
        
        # 配置各种文本标签
        self.gpt_history.tag_configure("code_python", font=("Courier New", 12), background="#f8f8f8", foreground="#000000", spacing1=5, spacing3=5)
        self.gpt_history.tag_configure("code_javascript", font=("Courier New", 12), background="#f8f8f8", foreground="#000000", spacing1=5, spacing3=5)
        self.gpt_history.tag_configure("code_html", font=("Courier New", 12), background="#f8f8f8", foreground="#000000", spacing1=5, spacing3=5)
        self.gpt_history.tag_configure("code_css", font=("Courier New", 12), background="#f8f8f8", foreground="#000000", spacing1=5, spacing3=5)
        self.gpt_history.tag_configure("code_default", font=("Courier New", 12), background="#f8f8f8", foreground="#000000", spacing1=5, spacing3=5)
        self.gpt_history.tag_configure("code_inline", font=("Courier New", 12), background="#f0f0f0", spacing1=2, spacing3=2)
        self.gpt_history.tag_configure("bold", font=("Arial", 14, "bold"))
        self.gpt_history.tag_configure("italic", font=("Arial", 14, "italic"))
        self.gpt_history.tag_configure("link", foreground="blue", underline=1)
        self.gpt_history.tag_configure("list_item", lmargin1=20, lmargin2=20, spacing1=2, spacing3=2)
        self.gpt_history.tag_configure("table_cell", font=("Arial", 12), background="#f8f8f8", spacing1=2, spacing3=2)
        self.gpt_history.tag_configure("heading1", font=("Arial", 20, "bold"), spacing1=10, spacing3=5)
        self.gpt_history.tag_configure("heading2", font=("Arial", 18, "bold"), spacing1=8, spacing3=4)
        self.gpt_history.tag_configure("heading3", font=("Arial", 16, "bold"), spacing1=6, spacing3=3)
        self.gpt_history.tag_configure("blockquote", font=("Arial", 14), background="#f8f8f8", lmargin1=20, lmargin2=20, spacing1=5, spacing3=5)
        
        gpt_scroll = ttk.Scrollbar(gpt_hist_content_frame, orient="vertical", command=self.gpt_history.yview)
        gpt_scroll.grid(row=0, column=1, sticky="ns")
        self.gpt_history['yscrollcommand'] = gpt_scroll.set

        # 添加到 PanedWindow，设置 minsize
        history_paned.add(user_hist_outer, stretch="always", minsize=120)
        history_paned.add(gpt_hist_outer, stretch="always", minsize=120)

        # 让 PanedWindow 两侧内容自适应拉伸
        user_hist_outer.rowconfigure(1, weight=1)
        user_hist_outer.columnconfigure(0, weight=1)
        gpt_hist_outer.rowconfigure(1, weight=1)
        gpt_hist_outer.columnconfigure(0, weight=1)
        user_hist_content_frame.rowconfigure(0, weight=1)
        user_hist_content_frame.columnconfigure(0, weight=1)
        gpt_hist_content_frame.rowconfigure(0, weight=1)
        gpt_hist_content_frame.columnconfigure(0, weight=1)

        # 保存按钮区
        save_frame = ttk.Frame(chat_frame)
        save_frame.grid(row=3, column=0, columnspan=2, pady=5, sticky="ew")
        save_frame.columnconfigure(0, weight=1)
        save_frame.columnconfigure(1, weight=1)
        save_frame.columnconfigure(2, weight=1)
        save_frame.columnconfigure(3, weight=1)
        
        # 第一排按钮：全部保存
        ttk.Button(save_frame, text="保存全部对话", command=self.save_all_turns_txt).grid(row=0, column=0, padx=5, pady=2)
        ttk.Button(save_frame, text="保存全部Word", command=self.save_all_turns_word).grid(row=0, column=1, padx=5, pady=2)
        ttk.Button(save_frame, text="保存全部HTML", command=self.save_all_turns_html).grid(row=0, column=2, padx=5, pady=2)
        ttk.Button(save_frame, text="保存历史对话", command=self.save_history_txt).grid(row=0, column=3, padx=5, pady=2)
        
        # 第二排按钮：单轮保存
        ttk.Button(save_frame, text="保存本轮对话", command=self.save_last_turn_txt).grid(row=1, column=0, padx=5, pady=2)
        ttk.Button(save_frame, text="保存本轮Word", command=self.save_last_turn_word).grid(row=1, column=1, padx=5, pady=2)
        ttk.Button(save_frame, text="保存本轮HTML", command=self.save_last_turn_html).grid(row=1, column=2, padx=5, pady=2)
        ttk.Button(save_frame, text="保存历史Word", command=self.save_history_word).grid(row=1, column=3, padx=5, pady=2)

        # 右键菜单
        self.create_history_context_menu()

    def toggle_browser(self):
        if self.driver is None:
            self.start_browser()
        else:
            self.stop_browser()

    def start_browser(self):
        try:
            options = uc.ChromeOptions()
            driver = uc.Chrome(version_main=136, options=options)
            self.driver = driver
            self.driver.get("https://chat.openai.com/")
            self.driver.maximize_window()
            self.start_button.configure(text="停止浏览器")
            self.status_var.set("浏览器已启动，请登录")

            # 登录检测（检查 class 包含 composer-btn 的 button）
            try:
                WebDriverWait(self.driver, 20).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, "button.composer-btn"))
                )
                self.status_var.set("已经登录成功")
            except TimeoutException:
                # 打印所有button的class和aria-label，便于调试
                buttons = self.driver.find_elements(By.TAG_NAME, "button")
                print("页面所有button的class和aria-label：")
                for btn in buttons:
                    print(f"class={btn.get_attribute('class')}, aria-label={btn.get_attribute('aria-label')}")
                # 未检测到，保持"请登录"状态
                pass

        except Exception as e:
            messagebox.showerror("错误", f"启动浏览器失败: {e}")

    def stop_browser(self):
        if self.driver:
            self.driver.quit()
            self.driver = None
            self.start_button.configure(text="启动浏览器")
            self.status_var.set("浏览器已关闭")

    def send_message(self):
        if not self.driver:
            messagebox.showwarning("警告", "请先启动浏览器！")
            return
        message = self.message_text.get("1.0", tk.END).strip()
        if not message:
            messagebox.showwarning("警告", "请输入消息！")
            return
        try:
            self.status_var.set("正在发送消息...")
            if self.send_message_to_chatgpt(message):
                response = self.get_last_response(self.is_first_message)
                if response:
                    self.is_first_message = False
                    self.status_var.set("消息发送成功")
                    self.turn_data.append({"question": message, "answer": response})
                    self.append_history(message, response)
                else:
                    self.status_var.set("获取回复失败")
            else:
                self.status_var.set("发送消息失败")
        except Exception as e:
            self.status_var.set(f"发生错误: {str(e)}")

    def append_history(self, question, answer):
        # 输入内容追加到左侧历史框，回复追加到右侧历史框，均保留全部内容
        self.user_history.config(state=tk.NORMAL)
        self.gpt_history.config(state=tk.NORMAL)
        
        # 添加分割线
        separator = "\n" + "="*50 + "\n"
        
        if self.user_history.index("end-1c") != "1.0":
            self.user_history.insert(tk.END, separator)
        if self.gpt_history.index("end-1c") != "1.0":
            self.gpt_history.insert(tk.END, separator)
            
        # 处理用户问题
        if question.strip():  # 只有当问题不为空时才添加
            self.user_history.insert(tk.END, question)
        
        # 处理GPT回复的富文本格式
        try:
            # 使用html2text转换HTML为Markdown
            h = html2text.HTML2Text()
            h.ignore_links = False
            h.ignore_images = False
            h.ignore_tables = False
            h.ignore_emphasis = False
            h.body_width = 0  # 不自动换行
            h.unicode_snob = True  # 使用Unicode字符
            h.ul_item_mark = "•"  # 使用圆点作为无序列表标记
            h.emphasis_mark = "*"  # 使用星号作为强调标记
            
            # 转换HTML为Markdown
            markdown_text = h.handle(answer)
            
            # 插入处理后的内容
            start_index = self.gpt_history.index("end-1c")
            self.gpt_history.insert(tk.END, markdown_text)
            
            # 应用样式
            self.apply_rich_text_styles(start_index)
        except Exception as e:
            print(f"处理富文本时出错: {str(e)}")
            # 如果处理失败，使用普通文本
            if answer.strip():  # 只有当回答不为空时才添加
                self.gpt_history.insert(tk.END, answer)
        
        self.user_history.see(tk.END)
        self.gpt_history.see(tk.END)
        self.user_history.config(state=tk.DISABLED)
        
        # 清空输入框
        self.message_text.delete("1.0", tk.END)

    def apply_rich_text_styles(self, start_index):
        """应用富文本样式"""
        try:
            # 获取插入的内容
            content = self.gpt_history.get(start_index, "end-1c")
            
            # 处理代码块
            code_pattern = r'```(\w*)\n(.*?)```'
            for match in re.finditer(code_pattern, content, re.DOTALL):
                lang = match.group(1) or 'default'
                code = match.group(2)
                code_start = self.gpt_history.search(code, start_index, "end-1c")
                if code_start:
                    code_end = f"{code_start}+{len(code)}c"
                    self.gpt_history.tag_add(f"code_{lang}", code_start, code_end)
            
            # 处理行内代码
            inline_code_pattern = r'`([^`]+)`'
            for match in re.finditer(inline_code_pattern, content):
                code = match.group(1)
                code_start = self.gpt_history.search(code, start_index, "end-1c")
                if code_start:
                    code_end = f"{code_start}+{len(code)}c"
                    self.gpt_history.tag_add("code_inline", code_start, code_end)
            
            # 处理粗体
            bold_pattern = r'\*\*(.*?)\*\*'
            for match in re.finditer(bold_pattern, content):
                text = match.group(1)
                text_start = self.gpt_history.search(text, start_index, "end-1c")
                if text_start:
                    text_end = f"{text_start}+{len(text)}c"
                    self.gpt_history.tag_add("bold", text_start, text_end)
            
            # 处理斜体
            italic_pattern = r'\*(.*?)\*'
            for match in re.finditer(italic_pattern, content):
                text = match.group(1)
                text_start = self.gpt_history.search(text, start_index, "end-1c")
                if text_start:
                    text_end = f"{text_start}+{len(text)}c"
                    self.gpt_history.tag_add("italic", text_start, text_end)
            
            # 处理链接
            link_pattern = r'\[(.*?)\]\((.*?)\)'
            for match in re.finditer(link_pattern, content):
                text = match.group(1)
                url = match.group(2)
                text_start = self.gpt_history.search(text, start_index, "end-1c")
                if text_start:
                    text_end = f"{text_start}+{len(text)}c"
                    self.gpt_history.tag_add("link", text_start, text_end)
                    # 存储链接URL
                    self.gpt_history.tag_bind("link", "<Button-1>", lambda e, url=url: self.open_url(url))
            
            # 处理列表
            list_pattern = r'^[\s]*[-*+]\s+(.*?)$'
            for match in re.finditer(list_pattern, content, re.MULTILINE):
                text = match.group(1)
                text_start = self.gpt_history.search(text, start_index, "end-1c")
                if text_start:
                    text_end = f"{text_start}+{len(text)}c"
                    self.gpt_history.tag_add("list_item", text_start, text_end)
            
            # 处理表格
            table_pattern = r'\|(.*?)\|'
            for match in re.finditer(table_pattern, content):
                text = match.group(1)
                text_start = self.gpt_history.search(text, start_index, "end-1c")
                if text_start:
                    text_end = f"{text_start}+{len(text)}c"
                    self.gpt_history.tag_add("table_cell", text_start, text_end)
        
        except Exception as e:
            print(f"应用样式时出错: {str(e)}")

    def open_url(self, url):
        """打开链接"""
        import webbrowser
        webbrowser.open(url)

    def create_history_context_menu(self):
        # 右键菜单
        self.hist_menu = tk.Menu(self.root, tearoff=0)
        self.hist_menu.add_command(label="复制", command=self.copy_selected_history)
        # 绑定事件
        self.user_history.bind("<Button-3>", self.show_user_hist_menu)
        self.gpt_history.bind("<Button-3>", self.show_gpt_hist_menu)
        # Mac下Ctrl+Button-1
        self.user_history.bind("<Control-Button-1>", self.show_user_hist_menu)
        self.gpt_history.bind("<Control-Button-1>", self.show_gpt_hist_menu)

    def show_user_hist_menu(self, event):
        try:
            self._hist_menu_target = self.user_history
            self.hist_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.hist_menu.grab_release()

    def show_gpt_hist_menu(self, event):
        try:
            self._hist_menu_target = self.gpt_history
            self.hist_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.hist_menu.grab_release()

    def copy_selected_history(self):
        widget = getattr(self, "_hist_menu_target", None)
        if widget:
            try:
                text = widget.get(tk.SEL_FIRST, tk.SEL_LAST)
                self.root.clipboard_clear()
                self.root.clipboard_append(text)
            except tk.TclError:
                pass

    def send_message_to_chatgpt(self, message):
        """发送消息到ChatGPT"""
        try:
            # 等待输入框出现
            selectors = [
                "p[data-placeholder='询问任何问题']",
                "p[data-placeholder='Ask anything']",
                "div[contenteditable='true']",
                "div[role='textbox']"
            ]
            
            textarea = None
            for selector in selectors:
                textarea = self.wait_for_element(By.CSS_SELECTOR, selector)
                if textarea:
                    break
            
            if textarea:
                # 使用 JavaScript 来设置内容
                self.driver.execute_script("arguments[0].innerHTML = arguments[1]", textarea, message)
                time.sleep(1)
                
                # 等待发送按钮出现并点击
                send_button = self.wait_for_element(By.CSS_SELECTOR, "button[data-testid='send-button'], button[aria-label='发送消息']")
                if send_button:
                    send_button.click()
                    return True
            return False
        except Exception as e:
            print(f"发送消息时出错: {str(e)}")
            return False

    def wait_for_element(self, by, value, timeout=10):
        """等待元素出现并返回"""
        try:
            element = WebDriverWait(self.driver, timeout).until(
                EC.presence_of_element_located((by, value))
            )
            return element
        except TimeoutException:
            return None

    def get_last_response(self, is_first_message=False):
        """等待最新回复完成（监控复制按钮和文本内容变化）"""
        try:
            max_wait_time = 60  # 最多等待60秒
            start_time = time.time()
            last_text = ""
            while time.time() - start_time < max_wait_time:
                try:
                    # 获取所有对话轮次
                    articles = self.driver.find_elements(By.CSS_SELECTOR, "article[data-testid^='conversation-turn-']")
                    if not articles:
                        time.sleep(1)
                        continue
                    last_article = articles[-1]
                    # 检查复制按钮
                    copy_button = None
                    try:
                        copy_button = last_article.find_element(By.CSS_SELECTOR, "button[data-testid='copy-turn-action-button']")
                    except Exception:
                        copy_button = None
                    # 检查AI回复文本
                    markdown_div = None
                    try:
                        markdown_div = last_article.find_element(By.CSS_SELECTOR, "div.markdown.prose")
                    except Exception:
                        markdown_div = None
                    # 同时满足：按钮可见+文本有新变化
                    if copy_button and copy_button.is_displayed() and markdown_div:
                        # 获取innerHTML而不是text
                        current_html = markdown_div.get_attribute("innerHTML")
                        if current_html and current_html != last_text:
                            return current_html
                        last_text = current_html
                    time.sleep(1)
                except Exception as e:
                    print(f"查找回复时出错: {str(e)}")
                    time.sleep(1)
            return None
        except Exception as e:
            print(f"获取回复时出错: {str(e)}")
            return None

    def wait_for_any_copy_button(self, timeout=30):
        """等待任意复制按钮出现"""
        try:
            # 尝试多个可能的复制按钮选择器
            selectors = [
                "button[data-testid='copy-turn-action-button']",
                "button[aria-label='Copy message']",
                "button[aria-label='复制消息']",
                "button.copy-button"
            ]
            
            for selector in selectors:
                try:
                    WebDriverWait(self.driver, timeout/len(selectors)).until(
                        EC.presence_of_element_located((By.CSS_SELECTOR, selector))
                    )
                    return True
                except TimeoutException:
                    continue
                    
            return False
        except Exception as e:
            print(f"等待复制按钮时出错: {str(e)}")
            return False

    # 已移除保存JSON方法

    def _get_increment_filename(self, base, ext):
        """生成不重复的文件名（如 base.ext, base_1.ext, base_2.ext...）"""
        fname = f"{base}.{ext}"
        if not os.path.exists(fname):
            return fname
        idx = 1
        while True:
            fname = f"{base}_{idx}.{ext}"
            if not os.path.exists(fname):
                return fname
            idx += 1

    def save_all_turns_txt(self):
        if not self.turn_data:
            messagebox.showwarning("警告", "没有轮次数据可保存。")
            return
        # 询问保存模式
        mode = messagebox.askquestion("保存模式", "仅保存回答内容？\n选择'否'将保存问题+回答。")
        try:
            fname = self._get_increment_filename("chatgpt_all_turns", "txt")
            with open(fname, "w", encoding="utf-8") as f:
                for i, turn in enumerate(self.turn_data, 1):
                    if mode == "yes":
                        f.write(f"{turn['answer']}\n")
                        if i < len(self.turn_data):  # 如果不是最后一个回答，添加分割线
                            f.write("="*50 + "\n")
                    else:
                        f.write(f"Q: {turn['question']}\nA: {turn['answer']}\n\n")
            messagebox.showinfo("成功", f"已保存为 {os.path.basename(fname)}")
        except Exception as e:
            messagebox.showerror("错误", f"保存失败: {e}")

    def save_all_turns_word(self):
        if not self.turn_data:
            messagebox.showwarning("警告", "没有轮次数据可保存。")
            return
        # 询问保存模式
        mode = messagebox.askquestion("保存模式", "仅保存回答内容？\n选择'否'将保存问题+回答。")
        try:
            fname = self._get_increment_filename("chatgpt_all_turns", "docx")
            doc = Document()
            for i, turn in enumerate(self.turn_data, 1):
                if mode == "yes":
                    doc.add_paragraph(turn['answer'])
                    if i < len(self.turn_data):  # 如果不是最后一个回答，添加分割线
                        doc.add_paragraph("="*50)
                else:
                    doc.add_paragraph(f"Q: {turn['question']}")
                    doc.add_paragraph(f"A: {turn['answer']}")
                    doc.add_paragraph("")
            doc.save(fname)
            messagebox.showinfo("成功", f"已保存为 {os.path.basename(fname)}")
        except Exception as e:
            messagebox.showerror("错误", f"保存失败: {e}")

    # 已移除保存Markdown/HTML/PDF通用方法

    def save_all_turns_html(self):
        if not self.driver:
            messagebox.showwarning("警告", "请先启动浏览器并加载对话。")
            return
        # 询问保存模式
        mode = messagebox.askquestion("保存模式", "仅保存回答内容？\n选择'否'将保存问题+回答。")
        try:
            # 获取所有轮次的AI回复的 innerHTML
            articles = self.driver.find_elements(By.CSS_SELECTOR, "article[data-testid^='conversation-turn-']")
            html_blocks = []
            for idx, article in enumerate(articles, 1):
                try:
                    markdown_div = article.find_element(By.CSS_SELECTOR, "div.markdown.prose")
                    inner_html = markdown_div.get_attribute("innerHTML")
                    if mode == "yes":
                        html_blocks.append(f'<div class="ai-turn">{inner_html}</div>')
                        if idx < len(articles):  # 如果不是最后一个回答，添加分割线
                            html_blocks.append('<div class="separator">' + '='*50 + '</div>')
                    else:
                        user_div = article.find_element(By.CSS_SELECTOR, "div[data-message-author-role='user']")
                        user_html = user_div.get_attribute("innerHTML")
                        html_blocks.append(f'<div class="user-turn"><h3>Q:</h3>{user_html}</div>\n<div class="ai-turn"><h3>A:</h3>{inner_html}</div>')
                except Exception:
                    continue
            if not html_blocks:
                messagebox.showwarning("警告", "未能获取到AI回复内容。")
                return
            # 合并内容
            html_content = "\n".join(html_blocks)
            # 包裹完整HTML
            full_html = (
                '<html><head><meta charset="utf-8">'
                '<style>'
                'body { font-family: "Microsoft YaHei", Arial, sans-serif; background: #fff; color: #222;}'
                '.markdown.prose { font-size: 1em; }'
                'table { border-collapse: collapse; }'
                'td, th { border: 1px solid #ccc; padding: 4px 8px; }'
                'img { max-width: 100%; height: auto; }'
                '.separator { color: #ccc; margin: 1em 0; }'
                '</style></head><body>'
                f'{html_content}'
                '</body></html>'
            )
            fname = self._get_increment_filename("chatgpt_all_turns", "html")
            with open(fname, "w", encoding="utf-8") as f:
                f.write(full_html)
            messagebox.showinfo("成功", f"已保存为 {os.path.basename(fname)}")
        except Exception as e:
            messagebox.showerror("错误", f"保存失败: {e}")

    def save_last_turn_txt(self):
        if not self.turn_data:
            messagebox.showwarning("警告", "没有轮次数据可保存。")
            return
        # 询问保存模式
        mode = messagebox.askquestion("保存模式", "仅保存回答内容？\n选择'否'将保存问题+回答。")
        try:
            fname = self._get_increment_filename("chatgpt_last_turn", "txt")
            with open(fname, "w", encoding="utf-8") as f:
                turn = self.turn_data[-1]
                if mode == "yes":
                    f.write(f"{turn['answer']}\n")
                else:
                    f.write(f"Q: {turn['question']}\nA: {turn['answer']}\n")
            messagebox.showinfo("成功", f"已保存为 {os.path.basename(fname)}")
        except Exception as e:
            messagebox.showerror("错误", f"保存失败: {e}")

    def save_last_turn_word(self):
        if not self.turn_data:
            messagebox.showwarning("警告", "没有轮次数据可保存。")
            return
        # 询问保存模式
        mode = messagebox.askquestion("保存模式", "仅保存回答内容？\n选择'否'将保存问题+回答。")
        try:
            fname = self._get_increment_filename("chatgpt_last_turn", "docx")
            doc = Document()
            turn = self.turn_data[-1]
            if mode == "yes":
                doc.add_paragraph(turn['answer'])
            else:
                doc.add_paragraph(f"Q: {turn['question']}")
                doc.add_paragraph(f"A: {turn['answer']}")
            doc.save(fname)
            messagebox.showinfo("成功", f"已保存为 {os.path.basename(fname)}")
        except Exception as e:
            messagebox.showerror("错误", f"保存失败: {e}")

    def save_last_turn_html(self):
        if not self.driver:
            messagebox.showwarning("警告", "请先启动浏览器并加载对话。")
            return
        # 询问保存模式
        mode = messagebox.askquestion("保存模式", "仅保存回答内容？\n选择'否'将保存问题+回答。")
        try:
            # 获取最后一轮AI回复的 innerHTML
            articles = self.driver.find_elements(By.CSS_SELECTOR, "article[data-testid^='conversation-turn-']")
            if not articles:
                messagebox.showwarning("警告", "未能获取到AI回复内容。")
                return
            last_article = articles[-1]
            try:
                markdown_div = last_article.find_element(By.CSS_SELECTOR, "div.markdown.prose")
                inner_html = markdown_div.get_attribute("innerHTML")
                if mode == "yes":
                    html_content = f'<div class="ai-turn">{inner_html}</div>\n'
                else:
                    user_div = last_article.find_element(By.CSS_SELECTOR, "div[data-message-author-role='user']")
                    user_html = user_div.get_attribute("innerHTML")
                    html_content = f'<div class="user-turn"><h3>Q:</h3>{user_html}</div>\n<div class="ai-turn"><h3>A:</h3>{inner_html}</div>\n'
            except Exception:
                messagebox.showwarning("警告", "未能获取到AI回复内容。")
                return
            # 包裹完整HTML
            full_html = (
                '<html><head><meta charset="utf-8">'
                '<style>'
                'body { font-family: "Microsoft YaHei", Arial, sans-serif; background: #fff; color: #222;}'
                '.markdown.prose { font-size: 1em; }'
                'table { border-collapse: collapse; }'
                'td, th { border: 1px solid #ccc; padding: 4px 8px; }'
                'img { max-width: 100%; height: auto; }'
                '</style></head><body>'
                f'{html_content}'
                '</body></html>'
            )
            fname = self._get_increment_filename("chatgpt_last_turn", "html")
            with open(fname, "w", encoding="utf-8") as f:
                f.write(full_html)
            messagebox.showinfo("成功", f"已保存为 {os.path.basename(fname)}")
        except Exception as e:
            messagebox.showerror("错误", f"保存失败: {e}")

    def save_history_txt(self):
        if not self.driver:
            messagebox.showwarning("警告", "请先启动浏览器并加载对话。")
            return
        # 询问保存模式
        mode = messagebox.askquestion("保存模式", "仅保存回答内容？\n选择'否'将保存问题+回答。")
        try:
            articles = self.driver.find_elements(By.CSS_SELECTOR, "article[data-testid^='conversation-turn-']")
            qa_list = []
            for art in articles:
                try:
                    user_div = art.find_element(By.CSS_SELECTOR, "div[data-message-author-role='user']")
                    q = user_div.text.strip()
                except Exception:
                    q = ""
                try:
                    ai_div = art.find_element(By.CSS_SELECTOR, "div.markdown.prose")
                    a = ai_div.text.strip()
                except Exception:
                    a = ""
                if q or a:
                    qa_list.append((q, a))
            if not qa_list:
                messagebox.showwarning("警告", "未能获取到历史内容。")
                return
            fname = self._get_increment_filename("chatgpt_history", "txt")
            with open(fname, "w", encoding="utf-8") as f:
                for idx, (q, a) in enumerate(qa_list, 1):
                    if mode == "yes":
                        f.write(f"{a}\n")
                        if idx < len(qa_list):  # 如果不是最后一个回答，添加分割线
                            f.write("="*50 + "\n")
                    else:
                        if q:  # 只有当问题不为空时才写入
                            f.write(f"Q: {q}\n")
                        if a:  # 只有当回答不为空时才写入
                            f.write(f"A: {a}\n\n")
            messagebox.showinfo("成功", f"已保存为 {os.path.basename(fname)}")
        except Exception as e:
            messagebox.showerror("错误", f"保存失败: {e}")

    def save_history_word(self):
        if not self.driver:
            messagebox.showwarning("警告", "请先启动浏览器并加载对话。")
            return
        # 询问保存模式
        mode = messagebox.askquestion("保存模式", "仅保存回答内容？\n选择'否'将保存问题+回答。")
        try:
            articles = self.driver.find_elements(By.CSS_SELECTOR, "article[data-testid^='conversation-turn-']")
            qa_list = []
            for art in articles:
                try:
                    user_div = art.find_element(By.CSS_SELECTOR, "div[data-message-author-role='user']")
                    q = user_div.text.strip()
                except Exception:
                    q = ""
                try:
                    ai_div = art.find_element(By.CSS_SELECTOR, "div.markdown.prose")
                    a = ai_div.text.strip()
                except Exception:
                    a = ""
                if q or a:
                    qa_list.append((q, a))
            if not qa_list:
                messagebox.showwarning("警告", "未能获取到历史内容。")
                return
            fname = self._get_increment_filename("chatgpt_history", "docx")
            doc = Document()
            for idx, (q, a) in enumerate(qa_list, 1):
                if mode == "yes":
                    doc.add_paragraph(a)
                    if idx < len(qa_list):  # 如果不是最后一个回答，添加分割线
                        doc.add_paragraph("="*50)
                else:
                    if q:  # 只有当问题不为空时才写入
                        doc.add_paragraph(f"Q: {q}")
                    if a:  # 只有当回答不为空时才写入
                        doc.add_paragraph(f"A: {a}")
                    doc.add_paragraph("")
            doc.save(fname)
            messagebox.showinfo("成功", f"已保存为 {os.path.basename(fname)}")
        except Exception as e:
            messagebox.showerror("错误", f"保存失败: {e}")

def main():
    root = tk.Tk()
    app = ChatGPTGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()