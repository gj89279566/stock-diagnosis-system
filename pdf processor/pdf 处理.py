import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LTChar, LTTextBox, LTTextLine, LTRect
from docx import Document
from collections import defaultdict

# 全局字体信息缓存
font_info = defaultdict(lambda: {"sample": "", "count": 0})

class DebugWindow:
    def __init__(self, parent):
        self.window = tk.Toplevel(parent)
        self.window.title("调试信息")
        self.window.geometry("800x600")
        
        # 创建文本框和滚动条
        self.text = tk.Text(self.window, wrap=tk.WORD)
        scrollbar = ttk.Scrollbar(self.window, orient="vertical", command=self.text.yview)
        self.text.configure(yscrollcommand=scrollbar.set)
        
        # 布局
        self.text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
    def add_text(self, text):
        self.text.insert(tk.END, text + "\n")
        self.text.see(tk.END)
        
    def clear(self):
        self.text.delete(1.0, tk.END)

# GUI 主程序
class PDFExtractorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF 字体提取器")
        self.pdf_path = ""
        self.font_check_vars = {}
        self.debug_window = None
        
        # 创建主框架
        main_frame = ttk.Frame(root, padding="10")
        main_frame.pack(fill="both", expand=True)
        
        # 按钮区域
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill="x", pady=5)
        
        ttk.Button(button_frame, text="选择 PDF 文件", command=self.select_pdf).pack(side="left", padx=5)
        ttk.Button(button_frame, text="分析字体", command=self.analyze_fonts).pack(side="left", padx=5)
        ttk.Button(button_frame, text="导出 Word", command=self.export_selected_fonts).pack(side="left", padx=5)
        ttk.Button(button_frame, text="保留格式导出", command=self.export_with_format).pack(side="left", padx=5)
        ttk.Button(button_frame, text="调试窗口", command=self.show_debug_window).pack(side="left", padx=5)
        
        # 创建字体列表框架
        self.font_frame = ttk.LabelFrame(main_frame, text="字体列表", padding="5")
        self.font_frame.pack(fill="both", expand=True, pady=5)
        
        # 创建字体列表
        self.font_list = ttk.Treeview(
            self.font_frame,
            columns=("checkbox", "size", "font", "color", "bg_color", "sample", "count"),
            show="headings"
        )
        self.font_list.heading("checkbox", text="")
        self.font_list.heading("size", text="字体大小")
        self.font_list.heading("font", text="字体名称")
        self.font_list.heading("color", text="颜色")
        self.font_list.heading("bg_color", text="底纹")
        self.font_list.heading("sample", text="示例文本")
        self.font_list.heading("count", text="出现次数")

        # 设置列宽和对齐
        self.font_list.column("checkbox", width=32, anchor="center")
        self.font_list.column("size", width=70, anchor="center")
        self.font_list.column("font", width=160, anchor="w")
        self.font_list.column("color", width=90, anchor="center")
        self.font_list.column("bg_color", width=90, anchor="center")
        self.font_list.column("sample", width=260, anchor="w")
        self.font_list.column("count", width=60, anchor="center")
        
        # 添加滚动条
        scrollbar = ttk.Scrollbar(self.font_frame, orient="vertical", command=self.font_list.yview)
        self.font_list.configure(yscrollcommand=scrollbar.set)
        self.font_list.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

    def get_color_name(self, color):
        """将PDF颜色值转换为可读的颜色名称"""
        if not color:
            return "黑色"
        
        try:
            # 处理不同格式的颜色值
            if isinstance(color, (list, tuple)):
                if len(color) == 3:  # RGB
                    r, g, b = color
                elif len(color) == 4:  # CMYK
                    c, m, y, k = color
                    # 简单的CMYK到RGB转换
                    r = 255 * (1 - c) * (1 - k)
                    g = 255 * (1 - m) * (1 - k)
                    b = 255 * (1 - y) * (1 - k)
                else:
                    return f"颜色值({color})"
            else:
                return f"颜色值({color})"
            
            # 将RGB值转换为0-1范围
            r, g, b = r/255, g/255, b/255
            
            # 简单的颜色映射
            if r > 0.9 and g < 0.1 and b < 0.1:
                return "红色"
            elif r < 0.1 and g > 0.9 and b < 0.1:
                return "绿色"
            elif r < 0.1 and g < 0.1 and b > 0.9:
                return "蓝色"
            elif r > 0.9 and g > 0.9 and b < 0.1:
                return "黄色"
            elif r > 0.9 and g < 0.1 and b > 0.9:
                return "紫色"
            elif r < 0.1 and g > 0.9 and b > 0.9:
                return "青色"
            elif r > 0.9 and g > 0.9 and b > 0.9:
                return "白色"
            elif r < 0.1 and g < 0.1 and b < 0.1:
                return "黑色"
            else:
                return f"RGB({int(r*255)},{int(g*255)},{int(b*255)})"
        except Exception as e:
            return f"颜色值({color})"

    def select_pdf(self):
        path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if path:
            self.pdf_path = path
            messagebox.showinfo("提示", f"选中文件：{os.path.basename(path)}")

    def analyze_fonts(self):
        if not self.pdf_path:
            messagebox.showerror("错误", "请先选择 PDF 文件。")
            return

        # 清空现有数据
        font_info.clear()
        self.font_check_vars.clear()
        for item in self.font_list.get_children():
            self.font_list.delete(item)

        for page_layout in extract_pages(self.pdf_path):
            rects = []
            # 收集所有底纹矩形
            for element in page_layout:
                if isinstance(element, LTRect):
                    if hasattr(element, "non_stroking_color") and element.non_stroking_color:
                        rects.append({
                            "bbox": element.bbox,
                            "color": self.get_color_name(element.non_stroking_color)
                        })
            # 处理文本
            for element in page_layout:
                if isinstance(element, LTTextBox):
                    for text_line in element:
                        if isinstance(text_line, LTTextLine):
                            for char in text_line:
                                if isinstance(char, LTChar):
                                    size = round(char.size, 1)
                                    fontname = char.fontname
                                    color = char.graphicstate.ncolor
                                    color_name = self.get_color_name(color)
                                    # 判断底纹
                                    bg_color = "无"
                                    c_x0, c_y0, c_x1, c_y1 = char.bbox
                                    for rect in rects:
                                        r_x0, r_y0, r_x1, r_y1 = rect["bbox"]
                                        if (c_x0 >= r_x0 and c_x1 <= r_x1 and c_y0 >= r_y0 and c_y1 <= r_y1):
                                            bg_color = rect["color"]
                                            break
                                    # 使用组合键来区分不同的字体、颜色和底纹
                                    key = (size, fontname, color_name, bg_color)
                                    if len(font_info[key]["sample"]) < 30:
                                        font_info[key]["sample"] += char.get_text()
                                    font_info[key]["count"] += 1

        # 显示字体信息
        for (size, fontname, color_name, bg_color), info in sorted(font_info.items(), key=lambda x: (-x[0][0], x[0][1], x[0][2], x[0][3])):
            var = tk.BooleanVar(value=False)
            self.font_check_vars[(size, fontname, color_name, bg_color)] = var
            item = self.font_list.insert("", "end", values=(
                "□",
                f"{size}pt",
                fontname,
                color_name,
                bg_color,
                info["sample"].strip(),
                info["count"]
            ))
        self.font_list.bind('<ButtonRelease-1>', self.on_click)

    def on_click(self, event):
        item = self.font_list.identify_row(event.y)
        if item:
            column = self.font_list.identify_column(event.x)
            if column == "#1":
                values = self.font_list.item(item)['values']
                size = float(values[1].replace('pt', ''))
                font = values[2]
                color = values[3]
                bg_color = values[4]
                var = self.font_check_vars.get((size, font, color, bg_color))
                if var:
                    var.set(not var.get())
                    values = list(values)
                    values[0] = "☑" if var.get() else "□"
                    self.font_list.item(item, values=values)

    def debug_log(self, text):
        if self.debug_window:
            self.debug_window.add_text(text)

    def export_selected_fonts(self):
        selected_styles = []
        for item in self.font_list.get_children():
            values = self.font_list.item(item)['values']
            if values[0] == "☑":
                size = float(values[1].replace('pt', ''))
                font = values[2]
                color = values[3]
                bg_color = values[4]
                selected_styles.append((size, font, color, bg_color))
        if not selected_styles:
            messagebox.showerror("错误", "请先选择要导出的字体样式。")
            return
        if self.debug_window:
            self.debug_window.clear()
            self.debug_log("选中的样式：")
            for size, font, color, bg_color in selected_styles:
                self.debug_log(f"大小: {size}pt, 字体: {font}, 颜色: {color}, 底纹: {bg_color}")
            self.debug_log("\n开始处理文本...\n")
        doc = Document()
        current_paragraph = ""
        current_style = None
        last_position = None
        paragraph_break = False
        char_count = 0
        style_count = defaultdict(int)
        for page_layout in extract_pages(self.pdf_path):
            rects = []
            for element in page_layout:
                if isinstance(element, LTRect):
                    if hasattr(element, "non_stroking_color") and element.non_stroking_color:
                        rects.append({
                            "bbox": element.bbox,
                            "color": self.get_color_name(element.non_stroking_color)
                        })
            for element in page_layout:
                if isinstance(element, LTTextBox):
                    for text_line in element:
                        if isinstance(text_line, LTTextLine):
                            for char in text_line:
                                if isinstance(char, LTChar):
                                    size = round(char.size, 1)
                                    fontname = char.fontname
                                    color = char.graphicstate.ncolor
                                    color_name = self.get_color_name(color)
                                    bg_color = "无"
                                    c_x0, c_y0, c_x1, c_y1 = char.bbox
                                    for rect in rects:
                                        r_x0, r_y0, r_x1, r_y1 = rect["bbox"]
                                        if (c_x0 >= r_x0 and c_x1 <= r_x1 and c_y0 >= r_y0 and c_y1 <= r_y1):
                                            bg_color = rect["color"]
                                            break
                                    current_char_style = (size, fontname, color_name, bg_color)
                                    style_count[current_char_style] += 1
                                    if last_position is not None:
                                        if abs(char.bbox[1] - last_position) > 5:
                                            paragraph_break = True
                                    last_position = char.bbox[1]
                                    if current_char_style in selected_styles:
                                        char_count += 1
                                        if self.debug_window and char_count <= 100:
                                            self.debug_log(f"处理字符: '{char.get_text()}' - 样式: {current_char_style}")
                                        if paragraph_break or (current_style != current_char_style and current_paragraph):
                                            if current_paragraph:
                                                doc.add_paragraph(current_paragraph)
                                                if self.debug_window:
                                                    self.debug_log(f"添加段落: '{current_paragraph}'")
                                            current_paragraph = char.get_text()
                                            current_style = current_char_style
                                            paragraph_break = False
                                        else:
                                            current_paragraph += char.get_text()
            if current_paragraph:
                doc.add_paragraph(current_paragraph)
                if self.debug_window:
                    self.debug_log(f"添加最后段落: '{current_paragraph}'")
                current_paragraph = ""
                current_style = None
                paragraph_break = True
        if self.debug_window:
            self.debug_log("\n统计信息:")
            self.debug_log(f"总处理字符数: {char_count}")
            self.debug_log("\n选中样式的统计:")
            for style in selected_styles:
                count = style_count.get(style, 0)
                self.debug_log(f"样式 {style}: {count}个字符")
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word 文件", "*.docx")])
        if save_path:
            doc.save(save_path)
            messagebox.showinfo("成功", f"已导出 Word 文件：\n{save_path}")

    def export_with_format(self):
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_COLOR_INDEX
        from docx.oxml.ns import qn
        selected_styles = []
        for item in self.font_list.get_children():
            values = self.font_list.item(item)['values']
            if values[0] == "☑":
                size = float(values[1].replace('pt', ''))
                font = values[2]
                color = values[3]
                bg_color = values[4]
                selected_styles.append((size, font, color, bg_color))
        if not selected_styles:
            messagebox.showerror("错误", "请先选择要导出的字体样式。")
            return
        if self.debug_window:
            self.debug_window.clear()
            self.debug_log("选中的样式：")
            for size, font, color, bg_color in selected_styles:
                self.debug_log(f"大小: {size}pt, 字体: {font}, 颜色: {color}, 底纹: {bg_color}")
            self.debug_log("\n开始处理文本...\n")
        doc = Document()
        current_paragraph = doc.add_paragraph()
        current_style = None
        last_position = None
        paragraph_break = False
        char_count = 0
        style_count = defaultdict(int)
        def get_word_color(color_name):
            # 只处理RGB(,,)格式
            if color_name.startswith("RGB("):
                try:
                    rgb = color_name[4:-1].split(",")
                    return RGBColor(int(rgb[0]), int(rgb[1]), int(rgb[2]))
                except:
                    return None
            # 常见色名
            color_map = {"黑色": RGBColor(0,0,0), "白色": RGBColor(255,255,255), "红色": RGBColor(255,0,0), "绿色": RGBColor(0,255,0), "蓝色": RGBColor(0,0,255), "黄色": RGBColor(255,255,0)}
            return color_map.get(color_name, None)
        def get_word_highlight(bg_color):
            # 只支持有限的高亮色
            highlight_map = {
                "黄色": WD_COLOR_INDEX.YELLOW,
                "绿色": WD_COLOR_INDEX.BRIGHT_GREEN,
                "青色": WD_COLOR_INDEX.TURQUOISE,
                "红色": WD_COLOR_INDEX.PINK,
                "蓝色": WD_COLOR_INDEX.BLUE,
                "白色": WD_COLOR_INDEX.WHITE,
                "黑色": WD_COLOR_INDEX.BLACK
            }
            if bg_color in highlight_map:
                return highlight_map[bg_color]
            if bg_color.startswith("RGB("):
                # 近似映射
                rgb = tuple(int(x) for x in bg_color[4:-1].split(","))
                if rgb == (244,244,220):
                    return WD_COLOR_INDEX.YELLOW
                if rgb == (255,255,255):
                    return WD_COLOR_INDEX.WHITE
                if rgb == (0,0,0):
                    return WD_COLOR_INDEX.BLACK
            return None
        def is_bold(fontname):
            return "Bold" in fontname or "bold" in fontname
        def is_italic(fontname):
            return "Italic" in fontname or "Oblique" in fontname or "italic" in fontname or "oblique" in fontname
        def font_map(pdf_font):
            # 简单映射
            if "Times" in pdf_font:
                return "Times New Roman"
            if "Arial" in pdf_font:
                return "Arial"
            if "Hei" in pdf_font or "黑" in pdf_font:
                return "SimHei"
            if "Song" in pdf_font or "宋" in pdf_font:
                return "SimSun"
            return pdf_font
        for page_layout in extract_pages(self.pdf_path):
            rects = []
            for element in page_layout:
                if isinstance(element, LTRect):
                    if hasattr(element, "non_stroking_color") and element.non_stroking_color:
                        rects.append({
                            "bbox": element.bbox,
                            "color": self.get_color_name(element.non_stroking_color)
                        })
            for element in page_layout:
                if isinstance(element, LTTextBox):
                    for text_line in element:
                        if isinstance(text_line, LTTextLine):
                            for char in text_line:
                                if isinstance(char, LTChar):
                                    size = round(char.size, 1)
                                    fontname = char.fontname
                                    color = char.graphicstate.ncolor
                                    color_name = self.get_color_name(color)
                                    bg_color = "无"
                                    c_x0, c_y0, c_x1, c_y1 = char.bbox
                                    for rect in rects:
                                        r_x0, r_y0, r_x1, r_y1 = rect["bbox"]
                                        if (c_x0 >= r_x0 and c_x1 <= r_x1 and c_y0 >= r_y0 and c_y1 <= r_y1):
                                            bg_color = rect["color"]
                                            break
                                    current_char_style = (size, fontname, color_name, bg_color)
                                    style_count[current_char_style] += 1
                                    if last_position is not None:
                                        if abs(char.bbox[1] - last_position) > 5:
                                            paragraph_break = True
                                    last_position = char.bbox[1]
                                    if current_char_style in selected_styles:
                                        char_count += 1
                                        if self.debug_window and char_count <= 100:
                                            self.debug_log(f"处理字符: '{char.get_text()}' - 样式: {current_char_style}")
                                        if paragraph_break or (current_style != current_char_style and current_paragraph):
                                            if current_paragraph and len(current_paragraph.runs) > 0:
                                                current_paragraph = doc.add_paragraph()
                                            current_style = current_char_style
                                            paragraph_break = False
                                        run = current_paragraph.add_run(char.get_text())
                                        # 设置格式
                                        run.font.size = Pt(size)
                                        run.font.name = font_map(fontname)
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_map(fontname))
                                        word_color = get_word_color(color_name)
                                        if word_color:
                                            run.font.color.rgb = word_color
                                        run.bold = is_bold(fontname)
                                        run.italic = is_italic(fontname)
                                        highlight = get_word_highlight(bg_color)
                                        if highlight:
                                            run.font.highlight_color = highlight
                                # 段落结束
                            if paragraph_break:
                                current_paragraph = doc.add_paragraph()
                        # 文本框结束
            # 页面结束
        if self.debug_window:
            self.debug_log("\n统计信息:")
            self.debug_log(f"总处理字符数: {char_count}")
            self.debug_log("\n选中样式的统计:")
            for style in selected_styles:
                count = style_count.get(style, 0)
                self.debug_log(f"样式 {style}: {count}个字符")
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word 文件", "*.docx")])
        if save_path:
            doc.save(save_path)
            messagebox.showinfo("成功", f"已导出 Word 文件（保留格式）：\n{save_path}")

    def show_debug_window(self):
        if self.debug_window is None:
            self.debug_window = DebugWindow(self.root)
        else:
            self.debug_window.window.deiconify()

# 启动程序
if __name__ == "__main__":
    root = tk.Tk()
    app = PDFExtractorApp(root)
    root.mainloop()