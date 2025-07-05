import pdfplumber
from tkinter import filedialog, Tk, Button, Checkbutton, BooleanVar, Label, Frame, Canvas, Scrollbar
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from PIL import Image
import io
from collections import defaultdict

# 设置字体大小判断阈值
TITLE_SIZE_THRESHOLD = 14
SUBTITLE_SIZE_THRESHOLD = 12

def detect_text_style(text, font_size, font_name=None):
    """智能检测文本样式"""
    style = {
        'bold': False,
        'italic': False,
        'size': font_size,
        'color': RGBColor(0, 0, 0)  # 默认黑色
    }
    
    # 根据字体大小判断标题级别
    if font_size >= TITLE_SIZE_THRESHOLD:
        style['is_title'] = True
        style['level'] = 1
    elif font_size >= SUBTITLE_SIZE_THRESHOLD:
        style['is_title'] = True
        style['level'] = 2
    else:
        style['is_title'] = False
    
    return style

def apply_style_to_paragraph(paragraph, style):
    """应用样式到段落"""
    if style['is_title']:
        paragraph.style = f'Heading {style["level"]}'
    else:
        paragraph.style = 'Normal'
    
    # 设置字体大小
    for run in paragraph.runs:
        run.font.size = Pt(style['size'])
        run.font.color.rgb = style['color']
        if style.get('bold'):
            run.font.bold = True
        if style.get('italic'):
            run.font.italic = True

def process_table(table, doc):
    """处理表格并保持样式"""
    if not table:
        return
    
    # 创建表格
    table_doc = doc.add_table(rows=len(table), cols=len(table[0]))
    table_doc.style = 'Table Grid'
    
    # 填充表格内容并设置样式
    for i, row in enumerate(table):
        for j, cell in enumerate(row):
            if cell:
                cell_obj = table_doc.cell(i, j)
                cell_obj.text = cell
                # 设置单元格样式
                for paragraph in cell_obj.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def extract_images(page, doc):
    """提取并保存图片"""
    images = page.images
    for img in images:
        try:
            # 将图片数据转换为PIL Image对象
            image = Image.open(io.BytesIO(img['stream'].get_data()))
            # 保存图片到临时文件
            temp_path = f"temp_img_{hash(str(img))}.png"
            image.save(temp_path)
            # 添加图片到文档
            doc.add_picture(temp_path, width=Inches(6))
            # 删除临时文件
            os.remove(temp_path)
        except Exception as e:
            print(f"处理图片时出错: {e}")

def convert_pdf_to_word(pdf_path, keep_images, preserve_layout, handle_tables):
    doc = Document()
    BODY_SIZE_THRESHOLD = 12  # 设定正文最大字号

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            words = page.extract_words(keep_blank_chars=True, extra_attrs=['size'])
            # 按y坐标分组，模拟行
            lines = {}
            for word in words:
                if word['size'] <= BODY_SIZE_THRESHOLD:
                    y0 = round(word['top'])
                    if y0 not in lines:
                        lines[y0] = []
                    lines[y0].append(word['text'])
            # 按y坐标排序，逐行写入
            for y in sorted(lines.keys()):
                line_text = ' '.join(lines[y]).strip()
                if line_text:
                    doc.add_paragraph(line_text, style='Normal')

    output_path = pdf_path.replace('.pdf', '_body_only.docx')
    doc.save(output_path)
    print(f"仅正文文档已保存到：{output_path}")
    return output_path

# --- 字体扫描与排序 ---
def scan_pdf_fonts(pdf_path, max_examples=3):
    font_dict = defaultdict(list)  # {(fontname, size): [示例文本,...]}
    font_count = defaultdict(int)  # {(fontname, size): 总字数}
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            words = page.extract_words(keep_blank_chars=True, extra_attrs=['size', 'fontname'])
            for word in words:
                key = (word['fontname'], word['size'])
                if len(font_dict[key]) < max_examples:
                    font_dict[key].append(word['text'])
                font_count[key] += len(word['text'])
    font_list = []
    for (fontname, size), examples in font_dict.items():
        font_list.append({
            'fontname': fontname,
            'size': size,
            'examples': examples,
            'count': font_count[(fontname, size)]
        })
    font_list.sort(key=lambda x: x['count'], reverse=True)
    return font_list

# --- 全局变量 ---
selected_pdf_path = None
font_check_vars = []
font_options = []
font_checks_frame = None

# --- 步骤1：选择PDF并扫描字体 ---
def open_pdf_and_scan_fonts():
    global selected_pdf_path, font_check_vars, font_options, font_checks_frame
    file_path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
    if file_path:
        selected_pdf_path = file_path
        print(f"选择的文件是：{file_path}")
        font_options = scan_pdf_fonts(file_path)
        print("扫描到的字体组合：", font_options)
        if font_checks_frame:
            font_checks_frame.destroy()
        font_check_vars = []
        font_checks_frame = Frame(main_frame)
        font_checks_frame.pack(pady=10, fill='both', expand=True)
        if not font_options:
            Label(font_checks_frame, text="未检测到字体信息，请更换PDF文件。", fg='red').pack()
            return
        Label(font_checks_frame, text="请选择要提取的字体（字体名, 字号, 字数, 示例）：", font=('Arial', 10)).pack(anchor='w')
        # 滚动条
        canvas = Canvas(font_checks_frame, height=250)
        scrollbar = Scrollbar(font_checks_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = Frame(canvas)
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(
                scrollregion=canvas.bbox("all")
            )
        )
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        # 字体选项
        for idx, font in enumerate(font_options):
            var = BooleanVar()
            font_check_vars.append(var)
            example = ', '.join(font['examples'])
            text = f"{font['fontname']} 号:{font['size']} 字数:{font['count']} 示例: {example}"
            Checkbutton(scrollable_frame, text=text, variable=var, wraplength=400, anchor='w', justify='left').pack(anchor='w')

# --- 步骤2：根据用户选择生成Word ---
def execute_conversion_by_font():
    if not selected_pdf_path or not font_options:
        print("请先选择PDF并勾选字体！")
        return
    selected_fonts = set()
    for idx, var in enumerate(font_check_vars):
        if var.get():
            selected_fonts.add((font_options[idx]['fontname'], font_options[idx]['size']))
    if not selected_fonts:
        print("请至少选择一个字体！")
        return
    doc = Document()
    with pdfplumber.open(selected_pdf_path) as pdf:
        for page in pdf.pages:
            words = page.extract_words(keep_blank_chars=True, extra_attrs=['size', 'fontname'])
            lines = {}
            for word in words:
                key = (word['fontname'], word['size'])
                if key in selected_fonts:
                    y0 = round(word['top'])
                    if y0 not in lines:
                        lines[y0] = []
                    lines[y0].append(word['text'])
            for y in sorted(lines.keys()):
                line_text = ' '.join(lines[y]).strip()
                if line_text:
                    doc.add_paragraph(line_text, style='Normal')
    output_path = selected_pdf_path.replace('.pdf', '_selected_fonts.docx')
    doc.save(output_path)
    print(f"已根据选择的字体生成Word：{output_path}")

# --- GUI部分 ---
root = Tk()
root.title("PDF 到 Word 转换器")
root.geometry("600x600")

main_frame = Frame(root, padx=20, pady=20)
main_frame.pack(expand=True, fill='both')

Label(main_frame, text="1. 选择PDF文件", font=('Arial', 12)).pack(pady=5)
Button(main_frame, text="选择PDF并扫描字体", command=open_pdf_and_scan_fonts, bg='#2196F3', fg='white').pack(pady=5)

font_checks_frame = None

Label(main_frame, text="2. 勾选需要提取的字体后，点击下方按钮", font=('Arial', 12)).pack(pady=10)
Button(main_frame, text="开始转换（仅勾选字体）", command=execute_conversion_by_font, bg='#4CAF50', fg='white', font=('Arial', 12), padx=20, pady=10).pack(pady=20)

root.mainloop()