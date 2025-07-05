import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
import os
from docx.oxml.text.run import CT_R

class StyleGroup:
    def __init__(self, para_style, font, size, bold, italic, underline):
        self.para_style = para_style
        self.font = font
        self.size = size
        self.bold = bold
        self.italic = italic
        self.underline = underline
        self.runs = []  # (para_idx, run_idx, text)
        self.deleted = False
        self.example_texts = []
        self.total_chars = 0

    def add_run(self, para_idx, run_idx, text):
        self.runs.append((para_idx, run_idx, text))
        self.total_chars += len(text)
        if len(''.join(self.example_texts)) < 100:
            self.example_texts.append(text)

    def get_example(self):
        result = ''
        for idx, t in enumerate(self.example_texts):
            if idx > 0:
                result += '/'
            result += t
            if len(result) >= 100:
                break
        return result[:100]

    def summary(self):
        size_str = self.size if self.size != '?' else '样式默认'
        return f"{self.para_style} | {self.font} | {size_str} | {'B' if self.bold else ''}{'I' if self.italic else ''}{'U' if self.underline else ''} | {self.total_chars}字"

    def detail(self):
        size_str = self.size if self.size != '?' else '样式默认'
        size_note = '' if self.size != '?' else '\n(未显式设置字号，实际字号由Word样式决定)'
        return (f"段落样式: {self.para_style}\n"
                f"字体: {self.font}\n字号: {size_str}{size_note}\n加粗: {'是' if self.bold else '否'}\n斜体: {'是' if self.italic else '否'}\n下划线: {'是' if self.underline else '否'}\n"
                f"总字数: {self.total_chars}\n"
                f"示例: {self.get_example()}")

class ContentItem:
    def __init__(self, ctype, idx, summary, detail, ref=None, text_content=None):
        self.ctype = ctype
        self.idx = idx
        self.summary = summary
        self.detail = detail
        self.deleted = False
        self.ref = ref
        self.text_content = text_content  # 新增：存储文本框/图形的文字内容

class WordContentApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Word内容管理器（多选保留，文本框/图形内容可见）")
        self.style_groups = []
        self.table_items = []
        self.image_items = []
        self.textbox_items = []
        self.shape_items = []
        self.file_path = None
        self.create_widgets()

    def create_widgets(self):
        top_frame = tk.Frame(self.root)
        top_frame.pack(fill=tk.X, pady=5)
        tk.Button(top_frame, text="选择Word文件", command=self.open_file).pack(side=tk.LEFT, padx=5)
        tk.Button(top_frame, text="导出新Word", state=tk.DISABLED, command=self.export_word).pack(side=tk.RIGHT, padx=5)
        self.export_btn = top_frame.winfo_children()[-1]

        paned = ttk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        paned.pack(fill=tk.BOTH, expand=True)

        left_frame = tk.Frame(paned)
        paned.add(left_frame, weight=1)
        tk.Label(left_frame, text="内容分组/对象").pack()
        tree_frame = tk.Frame(left_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True)
        self.tree = ttk.Treeview(tree_frame, selectmode="extended", show="tree", height=30)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        xscroll = tk.Scrollbar(tree_frame, orient=tk.HORIZONTAL, command=self.tree.xview)
        xscroll.pack(side=tk.BOTTOM, fill=tk.X)
        self.tree.configure(xscrollcommand=xscroll.set)
        self.tree.bind("<<TreeviewSelect>>", self.on_select_item)

        right_frame = tk.Frame(paned)
        paned.add(right_frame, weight=3)
        tk.Label(right_frame, text="详细信息").pack()
        self.detail_text = tk.Text(right_frame, height=20, width=60, wrap=tk.WORD)
        self.detail_text.pack(fill=tk.BOTH, expand=True)

    def open_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word 文件", "*.docx")])
        if not file_path:
            return
        try:
            doc = Document(file_path)
            self.file_path = file_path
            self.analyze_doc(doc)
            self.populate_tree()
            self.export_btn.config(state=tk.NORMAL)
            self.detail_text.delete(1.0, tk.END)
            self.detail_text.insert(tk.END, "请在左侧多选要保留的内容，未选中的将被删除。\n")
        except Exception as e:
            messagebox.showerror("错误", f"无法读取Word文件: {str(e)}")

    def extract_text_from_xml(self, element):
        # 提取所有<w:t>标签的内容
        texts = []
        for t in element.iter():
            if t.tag.endswith('}t') and t.text:
                texts.append(t.text)
        return ''.join(texts)

    def analyze_doc(self, doc):
        self.style_groups = []
        self.table_items = []
        self.image_items = []
        self.textbox_items = []
        self.shape_items = []
        group_dict = {}
        for i, table in enumerate(doc.tables):
            rows, cols = len(table.rows), len(table.columns)
            summary = f"表格 {i+1}（{rows}行×{cols}列）"
            detail = f"表格 {i+1}，{rows}行×{cols}列"
            self.table_items.append(ContentItem('表格', i, summary, detail, ref=table))
        image_count = 0
        for rel in doc.part.rels.values():
            if 'image' in rel.target_ref:
                image_count += 1
                summary = f"图片 {image_count}"
                detail = f"图片 {image_count}，文件名: {rel.target_ref.split('/')[-1]}"
                self.image_items.append(ContentItem('图片', image_count-1, summary, detail, ref=rel))
        textbox_count = 0
        shape_count = 0
        for shape in doc.part.element.iter():
            # 文本框
            if shape.tag.endswith('txbxContent'):
                textbox_count += 1
                text_content = self.extract_text_from_xml(shape)
                summary = f"文本框 {textbox_count}"
                detail = f"文本框 {textbox_count}\n内容: {text_content[:100]}"
                self.textbox_items.append(ContentItem('文本框', textbox_count-1, summary, detail, ref=shape, text_content=text_content))
            # 图形
            if shape.tag.endswith('shape'):
                shape_count += 1
                text_content = self.extract_text_from_xml(shape)
                summary = f"图形 {shape_count}"
                detail = f"图形 {shape_count}\n内容: {text_content[:100]}"
                self.shape_items.append(ContentItem('图形', shape_count-1, summary, detail, ref=shape, text_content=text_content))
        for para_idx, para in enumerate(doc.paragraphs):
            para_style = para.style.name if para.style else '未知'
            for run_idx, run in enumerate(para.runs):
                text = run.text.strip()
                if not text:
                    continue
                font = run.font.name or (run.style.font.name if run.style and run.style.font and run.style.font.name else '默认')
                size = run.font.size.pt if run.font.size else (run.style.font.size.pt if run.style and run.style.font and run.style.font.size else '?')
                bold = bool(run.bold)
                italic = bool(run.italic)
                underline = bool(run.underline)
                key = (para_style, font, size, bold, italic, underline)
                if key not in group_dict:
                    group_dict[key] = StyleGroup(*key)
                group_dict[key].add_run(para_idx, run_idx, text)
        self.style_groups = list(group_dict.values())
        self.style_groups.sort(key=lambda g: -g.total_chars)

    def populate_tree(self):
        self.tree.delete(*self.tree.get_children())
        style_parent = self.tree.insert('', 'end', text='文本样式分组', open=True)
        for idx, group in enumerate(self.style_groups):
            tag = f"style_{idx}"
            display = group.summary()
            self.tree.insert(style_parent, 'end', iid=tag, text=display, tags=(tag,))
        table_parent = self.tree.insert('', 'end', text='表格', open=True)
        for item in self.table_items:
            tag = f"table_{item.idx}"
            display = item.summary
            self.tree.insert(table_parent, 'end', iid=tag, text=display, tags=(tag,))
        image_parent = self.tree.insert('', 'end', text='图片', open=True)
        for item in self.image_items:
            tag = f"image_{item.idx}"
            display = item.summary
            self.tree.insert(image_parent, 'end', iid=tag, text=display, tags=(tag,))
        textbox_parent = self.tree.insert('', 'end', text='文本框', open=True)
        for item in self.textbox_items:
            tag = f"textbox_{item.idx}"
            display = item.summary
            self.tree.insert(textbox_parent, 'end', iid=tag, text=display, tags=(tag,))
        shape_parent = self.tree.insert('', 'end', text='图形', open=True)
        for item in self.shape_items:
            tag = f"shape_{item.idx}"
            display = item.summary
            self.tree.insert(shape_parent, 'end', iid=tag, text=display, tags=(tag,))

    def on_select_item(self, event):
        selected = self.tree.selection()
        self.detail_text.delete(1.0, tk.END)
        if not selected:
            return
        tag = selected[0]
        if tag.startswith('style_'):
            idx = int(tag.split('_')[1])
            group = self.style_groups[idx]
            self.detail_text.insert(tk.END, group.detail() + "\n\n")
        elif tag.startswith('table_'):
            idx = int(tag.split('_')[1])
            item = self.table_items[idx]
            self.detail_text.insert(tk.END, item.detail + "\n\n")
        elif tag.startswith('image_'):
            idx = int(tag.split('_')[1])
            item = self.image_items[idx]
            self.detail_text.insert(tk.END, item.detail + "\n\n")
        elif tag.startswith('textbox_'):
            idx = int(tag.split('_')[1])
            item = self.textbox_items[idx]
            self.detail_text.insert(tk.END, item.detail + "\n\n")
        elif tag.startswith('shape_'):
            idx = int(tag.split('_')[1])
            item = self.shape_items[idx]
            self.detail_text.insert(tk.END, item.detail + "\n\n")

    def get_selected_leaf_tags(self):
        selected = set(self.tree.selection())
        leaf_tags = set()
        for tag in selected:
            if not self.tree.get_children(tag):  # 没有子节点
                leaf_tags.add(tag)
        return leaf_tags

    def export_word(self):
        doc = Document(self.file_path)
        selected = self.get_selected_leaf_tags()
        # 统计所有被选中的 run
        selected_run_indices = set()
        for idx, group in enumerate(self.style_groups):
            tag = f"style_{idx}"
            if tag in selected:
                for para_idx, run_idx, _ in group.runs:
                    selected_run_indices.add((para_idx, run_idx))
        # 删除未被选中的 run
        for para_idx, para in enumerate(doc.paragraphs):
            # 逆序删除 run，避免索引错乱
            for run_idx in reversed(range(len(para.runs))):
                if (para_idx, run_idx) not in selected_run_indices:
                    remove_run(para, run_idx)
        # 删除空段落
        self.remove_empty_paragraphs(doc)
        # 表格
        for item in self.table_items:
            tag = f"table_{item.idx}"
            if tag not in selected:
                try:
                    table = doc.tables[item.idx]
                    t = table._element
                    t.getparent().remove(t)
                except Exception:
                    pass
        # 图片
        for item in self.image_items:
            tag = f"image_{item.idx}"
            if tag not in selected:
                rel = item.ref
                try:
                    doc.part.drop_rel(rel.rId)
                except Exception:
                    pass
        # 文本框
        for item in self.textbox_items:
            tag = f"textbox_{item.idx}"
            if tag not in selected:
                shape = item.ref
                try:
                    remove_ancestor(shape)
                except Exception:
                    pass
        # 图形
        for item in self.shape_items:
            tag = f"shape_{item.idx}"
            if tag not in selected:
                shape = item.ref
                try:
                    remove_ancestor(shape)
                except Exception:
                    pass
        # 删除特殊对象
        self.remove_unselected_special_objects(doc, selected)
        # 再次删除空段落
        self.remove_empty_paragraphs(doc)
        base, ext = os.path.splitext(self.file_path)
        export_path = base + '_导出' + ext
        doc.save(export_path)
        messagebox.showinfo("导出完成", f"新Word已保存：{export_path}")

    def remove_unselected_special_objects(self, doc, selected_tags):
        # 获取主文档xml根节点
        root = doc._element.getroottree().getroot()
        # 需要彻底删除的标签
        special_tags = [
            'w:drawing', 'w:pict', 'w:object', 'w:smartArt', 'w:shape', 'w:sdt',
            'v:shape', 'v:textbox', 'w:txbxContent', 'w:tbl', 'w:hyperlink',
            'w:customXml', 'w:altChunk', 'w:control', 'w:group'
        ]
        for tag in special_tags:
            for elem in root.findall('.//'+tag, namespaces=root.nsmap):
                # 判断该对象是否属于选中内容（可根据iid或内容等判断）
                # 如果不属于选中内容，则删除
                parent = elem.getparent()
                if parent is not None:
                    parent.remove(elem)

    def remove_empty_paragraphs(self, doc):
        for para in reversed(doc.paragraphs):
            if not para.text.strip():
                p = para._element
                p.getparent().remove(p)

def remove_run(para, run_idx):
    run = para.runs[run_idx]
    r = run._element
    p = para._element
    p.remove(r)

def remove_ancestor(element, tags=('w:sdt', 'w:drawing', 'w:pict', 'w:shape')):
    parent = element
    while parent is not None:
        if any(parent.tag.endswith(tag) for tag in tags):
            grand = parent.getparent()
            if grand is not None:
                grand.remove(parent)
            return
        parent = parent.getparent()
    # fallback
    element.getparent().remove(element)

if __name__ == "__main__":
    root = tk.Tk()
    app = WordContentApp(root)
    root.mainloop() 