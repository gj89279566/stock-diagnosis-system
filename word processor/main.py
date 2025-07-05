


import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import re
import os

def clean_paragraphs(paragraphs):
    cleaned = []
    for para in paragraphs:
        text = para.text

        # 删除“横线+回车”或“回车+横线”的情况
        text = re.sub(r'^\s*[-–—]+\s*$', '', text)  # 删除单独成段的横线
        text = re.sub(r'[-–—]+\s*\n', '', text)     # 横线加回车
        text = re.sub(r'\n\s*[-–—]+', '', text)     # 回车加横线

        # 删除“ 空格 + 回车”
        text = re.sub(r' (\r?\n)+', ' ', text)

        cleaned.append(text.strip())

    # 删除连续空段
    final_cleaned = []
    skip = False
    for i, line in enumerate(cleaned):
        if line == '':
            if not skip:
                final_cleaned.append('')
                skip = True
        else:
            final_cleaned.append(line)
            skip = False

    return final_cleaned

def process_word_file(file_path, output_path):
    doc = Document(file_path)
    original_paragraphs = [p for p in doc.paragraphs]
    cleaned_texts = clean_paragraphs(original_paragraphs)

    new_doc = Document()
    for line in cleaned_texts:
        new_doc.add_paragraph(line)
    new_doc.save(output_path)

class WordCleanerApp:
    def __init__(self, master):
        self.master = master
        master.title("PDF转Word清理工具")

        self.label = tk.Label(master, text="请选择一个 Word 文件（.docx）进行清理")
        self.label.pack(pady=10)

        self.select_button = tk.Button(master, text="选择 Word 文件", command=self.select_file)
        self.select_button.pack(pady=5)

        self.process_button = tk.Button(master, text="开始处理", command=self.process_file, state=tk.DISABLED)
        self.process_button.pack(pady=5)

        self.file_path = None

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word 文档", "*.docx")])
        if file_path:
            self.file_path = file_path
            self.label.config(text=f"已选择：{os.path.basename(file_path)}")
            self.process_button.config(state=tk.NORMAL)

    def process_file(self):
        if not self.file_path:
            messagebox.showwarning("警告", "请先选择文件！")
            return
        save_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                 filetypes=[("Word 文档", "*.docx")],
                                                 title="保存清理后的文档为")
        if save_path:
            try:
                process_word_file(self.file_path, save_path)
                messagebox.showinfo("成功", f"处理完成，已保存为：\n{save_path}")
            except Exception as e:
                messagebox.showerror("错误", f"处理失败：{str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = WordCleanerApp(root)
    root.mainloop()